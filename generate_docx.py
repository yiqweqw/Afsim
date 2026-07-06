# -*- coding: utf-8 -*-
"""生成 externModel 参数分析报告 .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 辅助函数 ──
def heading(level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table(headers, rows):
    """添加带格式的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return table

def formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

# ============================================================
# 封面 / 标题
# ============================================================
title = doc.add_heading('AFSIM externModel 参数分析报告', level=0)
for run in title.runs:
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

para('雷达探测 · 侦察截获 · 电子干扰', bold=True)
para('分析日期：2026-06-16')
para('分析范围：externModel/sensor/ · externModel/ReconSensor/ · externModel/weapon/')
doc.add_page_break()

# ============================================================
# 0. 概述
# ============================================================
heading(1, '0. 概述')
para('本报告分析 AFSIM 2.9.0 中 externModel 目录下三个自定义模块的配置参数：')
para('• 模块一：RadarModel（雷达探测）—— externModel/sensor/')
para('• 模块二：WsfPassiveReconSensor（无源侦察截获）—— externModel/ReconSensor/')
para('• 模块三：PhaserWeapon / PhaserLethality（干扰与杀伤）—— externModel/weapon/')
para('每个模块分别列出所有影响计算结果的可加载参数，包含默认值、数学含义和在公式中的作用。')

# ============================================================
# 1. 雷达探测
# ============================================================
heading(1, '1. 雷达探测模块 (RadarModel)')

heading(2, '1.1 通过 applyConfig 传入的参数')
para('以下参数通过 std::map<std::string,double> 传入 RadarModel::applyConfig()：')
add_table(
    ['配置键', '物理量', '默认值', '在公式中的作用 / 含义'],
    [
        ['antennaGain', '天线增益 (dB)', '34.0', '收发天线峰值增益，影响天线方向图生成和雷达方程中的 Gt×Gr'],
        ['peakPower_', '峰值发射功率 (W)', '10000', '经占空比折算为平均功率，位于雷达方程分子'],
        ['fc', '工作频率 (Hz)', '9.7e9', '决定波长 λ=c/fc，雷达方程中的 λ² 项'],
        ['pulseWidth', '脉冲宽度 (s)', '8.0e-6', '与 PRF 共同计算占空比 dutyCycle = PW×PRF；脉压时参与 KC 计算'],
        ['bandwidth', '发射信号带宽 (Hz)', '5.0e6', '脉压增益 KC 计算因子；频率重合度分母'],
        ['Fn', '接收机噪声系数 (dB)', '3.0', 'Pn = k×T0×10^(Fn/10)×BW 中的噪声系数'],
        ['T0', '标准噪声温度 (K)', '290.0', '噪声功率计算中的参考温度'],
        ['L', '系统总损耗 (dB)', '5.0', '雷达方程分母 L_dB = 10^(L/10)'],
        ['pulseNum', '脉冲积累数', '128', '相干积累增益 KI=pulseNum；非相干 KI=√(pulseNum)'],
        ['Pfa', '虚警概率', '1.0e-6', 'Pd = 0.5×erfc(√(-ln(Pfa)) - √(SNR+0.5))'],
        ['Pd', '期望检测概率', '0.9', '反算最小可检测 SNR，用于最大作用距离估计'],
        ['rangeAccuracy', '距离测量精度 (m)', '50.0', '距离测量噪声标准差'],
        ['azimuthAccuracy', '方位测量精度 (deg)', '1.0', '方位/俯仰测量噪声标准差'],
        ['beamwidthAz_', '方位波束宽度 (deg)', '—(0)', 'isTargetInBeam() 波束覆盖判断'],
        ['beamwidthEl_', '俯仰波束宽度 (deg)', '—(0)', '俯仰方向波束覆盖判断'],
        ['bandwidthRx_', '接收机带宽 (Hz)', '自动派生', '有脉压时=bandwidth_；无脉压时=1/pulseWidth_'],
    ]
)

heading(2, '1.2 硬编码默认值（不通过 applyConfig 修改）')
add_table(
    ['参数', '值', '含义 / 作用'],
    [
        ['PRF_', '125000 Hz', '脉冲重复频率；不模糊距离 Ru = c/(2×PRF)'],
        ['N_', '40', '均匀线阵元数，决定方向图波束宽度和增益'],
        ['d_', '0.5', '阵元间距（波长归一化），决定栅瓣位置'],
        ['scanRate_', '10.0 deg/s', '波束机械扫描转速'],
        ['PCType_', '1（是）', '脉冲压缩标识：1→KC= PW×BW；0→KC=1'],
        ['coherentType_', '1（相干积累）', '积累类型：1→KI=pulseNum；0→KI=√(pulseNum)'],
        ['searchVolume_', '{0,360,0,30}', '{az_min, az_max, el_min, el_max}'],
    ]
)

heading(2, '1.3 关键雷达方程')
para('目标回波接收功率公式：', bold=True)
formula('Pr = KI × KC × P_avg × Gt × Gr × λ² × RCS / ((4π)³ × R⁴ × L_db)')

para('接收机噪声功率：', bold=True)
formula('Pn = k_Boltzmann × T0 × 10^(Fn/10) × BW_rx')

para('受干扰时的检测信噪比与探测概率：', bold=True)
formula('SNR = Pr / (Pn + Pjn)')
formula('Pd = 0.5 × erfc(√(-ln(Pfa)) - √(SNR + 0.5))')

para('受干扰最大探测距离：', bold=True)
formula('Rmax = [P_avg × Gt × Gr × λ² × RCS × KI × KC / ((4π)³ × (Pn+Pjn) × L)]^(1/4)')

doc.add_page_break()

# ============================================================
# 2. 侦察截获
# ============================================================
heading(1, '2. 无源侦察截获模块 (WsfPassiveReconSensor)')

heading(2, '2.1 可配置参数')
add_table(
    ['场景命令', '成员变量', '默认值', '含义'],
    [
        ['recon_sensitivity_db', 'mSensitivityDb', '-70.0 dBm', '接收机灵敏度门限；Pr_dBm ≥ 此值才触发检测'],
        ['recon_noise_t0', 'mNoiseT0K', '290.0 K', '接收机噪声温度'],
        ['recon_receiver_bandwidth', 'mEsmFallbackBandwidthHz', '500.0 MHz', '接收机分析带宽；决定噪声基底和频率重合度'],
        ['recon_snr_scale', 'mSnrScale', '5000.0', 'SNR 缩放因子；实际 SNR = Pr/(Pn) / scale'],
        ['recon_pulses_per_capture', 'mPulsesPerCapture', '32', '每次截获处理的脉冲数'],
        ['recon_process_time', 'mProcessTimeSec', '2.0 s', '帧处理时间间隔'],
        ['recon_pdw_num_limit', 'mPdwNumLimit', '1000', '帧缓存 PDW 上限；超限触发分选'],
        ['recon_rf_tolerance', 'mRfToleranceRatio', '0.01 (1%)', '分选时射频相对容差'],
        ['recon_pw_tolerance', 'mPwToleranceRatio', '0.10 (10%)', '分选时脉宽相对容差'],
        ['recon_doa_tolerance', 'mDoaToleranceRad', '~3°', '分选时到达角容差'],
        ['recon_cluster_max_distance', 'mClusterMaxDistance', '0.5', 'DBSCAN 聚类的最大特征距离'],
        ['recon_min_cluster_size', 'mMinClusterSize', '20', 'DBSCAN 最小聚类点数'],
        ['recon_sdif_neighbor_window', 'mSdifNeighborWindow', '50', 'SDIF 算法邻域窗大小'],
        ['recon_sdif_max_hist_bins', 'mSdifMaxHistogramBins', '1000', 'SDIF 直方图最大 bin 数'],
        ['recon_sdif_topk', 'mSdifTopK', '10', 'SDIF 保留的 TOP-K PRI 候选'],
        ['recon_identification_threshold', 'mIdentificationThreshold', '0.6', '雷达识别置信度门限；低于此输出 "Unknown"'],
        ['recon_use_native_context_builder', 'mUseNativeContextBuilder', 'true (1)', '是否回退到原生 ESM 链路构建电磁上下文'],
    ]
)

heading(2, '2.2 截获接收功率公式')
para('侦察接收的单程功率：', bold=True)
formula('Pr_esm = Pt_avg × Gt × Gr × λ² / ((4π)² × R² × Lr) × (BW_overlap / BW_radar)')
para('其中：', bold=False)
para('• Pt_avg = PeakPower × PulseWidth × PRF（从目标雷达发射机获取的平均功率）')
para('• Gt = 目标雷达天线峰值增益')
para('• Gr = 侦察接收天线峰值增益')
para('• Lr = 接收机内部损耗（线性值）')
para('• λ = c / 目标雷达频率')
para('• BW_overlap = 侦察带宽与雷达带宽的频率重合部分')

heading(2, '2.3 截获条件')
para('同时满足以下两条件：', bold=True)
para('1. 通视条件：斜距 ≤ 4.12 × (√H₁ + √H₂) × 1000 (m)')
para('2. 功率条件：Pr_esm ≥ 10^(灵敏度/10) × 1e-3 (W)，即 Pr_dBm ≥ mSensitivityDb')

heading(2, '2.4 测量误差模型')
para('各参数测量误差由 SNR 决定（所有误差均服从 N(0, σ²) 高斯分布）：')
add_table(
    ['测量量', '标准差 σ 公式'],
    [
        ['频率 σ_RF', '√3 / (π × PW × √SNR)'],
        ['到达角 σ_DOA', '√3 / (π × 10 × √SNR)'],
        ['脉宽 σ_PW', '√2 / (BW × √SNR)'],
        ['PRI σ_PRF', 'PRF / (√2 × π × √SNR)'],
        ['幅度 σ_PA', '√(P × 50) / √(4 × PW × BW × SNR)'],
    ]
)

doc.add_page_break()

# ============================================================
# 3. 干扰
# ============================================================
heading(1, '3. 电子干扰模块')

heading(2, '3.1 Jammer 数据结构参数（简化雷达模型）')
para('在自定义代码中通过 PhaserWeapon::SensorView 填充以下字段：')
add_table(
    ['字段', '来源', '含义'],
    [
        ['type', 'EA 技术类型', '"noise_jamming" 或 "false_target_jamming"'],
        ['fj', 'xmtr->GetFrequency()', '干扰机中心频率 (Hz)；决定干扰波长'],
        ['Bj', 'xmtr->GetBandwidth()', '干扰带宽 (Hz)；影响频率重合度'],
        ['Lj', 'xmtr->GetInternalLoss()', '干扰机内部损耗 (dB)'],
        ['EqPower', 'xmtr->GetPower()', '干扰发射功率 (W)'],
        ['antennaGain', 'xmtr->GetPeakAntennaGain()', '干扰天线峰值增益 (dB)'],
        ['FakeTargetNum', '硬编码 1', '假目标干扰时每次生成的假目标数量'],
        ['mBeamMinScanAz', '天线最小方位扫描限', '波束方位扫描下限'],
        ['mBeamMaxScanAz', '天线最大方位扫描限', '波束方位扫描上限'],
    ]
)

heading(2, '3.2 干扰接收功率方程（自定义简化模型）')
para('干扰机到达雷达接收机的功率：', bold=True)
formula('Pjn = KJC_radar × KJI_radar × EqPower × G_j2r × G_r2j × λj²')
formula('      / ((4π)² × R² × Lj) × (BW_overlap / Bj)')
para('其中：')
para('• KJC_radar = 被干扰雷达的脉冲压缩增益（从雷达获取）')
para('• KJI_radar = 被干扰雷达的积累增益（从雷达获取）')
para('• G_j2r      = 干扰天线在雷达方向上的增益')
para('• G_r2j      = 雷达天线在干扰方向上的增益')
para('• λj         = c / fj（干扰波长）')
para('• R          = 干扰机到雷达的斜距')
para('• BW_overlap = min(fju, Fru) - max(fjl, Frl)，即频率重合区间')

heading(2, '3.3 PhaserWeapon 完整参数（场景语法配置）')
para('通过 jammer_cssc_01.ag 语法的 PHASER_WEAPON 结构配置：')
add_table(
    ['场景命令', '成员变量', '默认值', '含义'],
    [
        ['maximum_spots_per_beam', 'JammerBeam::mMaxNumSpotsPerBeam', '—', '每波束最大干扰点数'],
        ['maximum_number_of_beams', 'JammerMode::mMaxNumBeams', '1', '最大波束数（1=阻塞式）'],
        ['maximum_number_of_spots', 'JammerMode::mMaximumNumSpots', '—', '总干扰点数上限'],
        ['frequency_following', 'JammerMode::mRepeater', 'false', '频率跟随（转发式干扰）'],
        ['signal_following', 'JammerMode::mRepeater', 'false', '信号跟随（转发式干扰）'],
        ['spot_power_distribution', 'mSpotPowerDistribution', 'average', '干扰点功率分配方式'],
        ['beam_power_distribution', 'mBeamPowerDistribution', 'average', '波束功率分配方式'],
        ['group_power_distribution', 'mGroupPowerDistribution', 'constant', '跨武器组功率分配方式'],
        ['jammer_group', 'mJammerGroupNameId', '—', '同一平台内干扰机组名称'],
    ]
)

heading(2, '3.4 发射机底层参数（XmtrRcvrAntenna 语法）')
para('通过标准 AFSIM 发射机配置：')
add_table(
    ['属性', 'Getter', '含义'],
    [
        ['frequency', 'GetFrequency()', '干扰中心频率 (Hz)'],
        ['bandwidth', 'GetBandwidth()', '干扰带宽 (Hz)'],
        ['power / peak_power', 'GetPower() / GetPeakPower()', '发射功率 (W)'],
        ['peak_antenna_gain', 'GetPeakAntennaGain()', '天线峰值增益 (dB)'],
        ['internal_loss', 'GetInternalLoss()', '内部损耗 (dB)'],
        ['azimuth_beamwidth', 'GetAzimuthBeamwidth()', '方位波束宽度 (deg)'],
        ['elevation_beamwidth', 'GetElevationBeamwidth()', '俯仰波束宽度 (deg)'],
    ]
)

heading(2, '3.5 杀伤效果参数（PhaserLethality）')
add_table(
    ['场景命令', '成员变量', '默认值', '含义'],
    [
        ['shield_damage_rate', 'mShieldDamageRate', '20.0', '每秒对护盾的伤害值'],
        ['armor_damage_rate', 'mArmorDamageRate', '20.0', '护盾归零后每秒对装甲的伤害值'],
    ]
)

para('杀伤逻辑：护盾 → 装甲 → 装甲归零 → 目标被消灭。')
para('护盾/装甲值从目标平台的 aux_data 中读取：phaser_shields / phaser_armor。')

heading(2, '3.6 完整 EW 效果链')
para('实际生效的干扰功率经过多层 EW 效果因子调制：', bold=True)
para('P_jam_effective  =  P_raw  ×  JammingPowerGain  ×  Blanking  ×  Cancelation',)
para('                    ×  ModulationGain  ×  JtoXGain  ×  RadiusFactor')
para('                    ×  RepeaterFactor  ×  RPJ_Gain')
para('这些因子分别对应 EA 技术效果：消隐(Blanking)、对消(Cancelation)、调制增益、')
para('干信比增益(JtoXGain)、距离因子、转发因子等，由 EW 效果框架(WsfEW_Effects)统一管理。')

doc.add_page_break()

# ============================================================
# 4. 参数影响权重
# ============================================================
heading(1, '4. 关键参数影响权重')

heading(2, '4.1 雷达探测')
add_table(
    ['排序', '参数', '影响程度', '说明'],
    [
        ['1', 'peakPower_ / averagePower_', '★★★★★', '位于雷达方程分子，探测距离 ∝ 四次方根'],
        ['2', 'antennaGain', '★★★★★', '收发双倍增益，Gt×Gr 为平方项'],
        ['3', 'fc → λ', '★★★★', '波长 λ² 项，频率越低探测距离越远'],
        ['4', 'pulseNum → KI_', '★★★★', '相参积累增益直乘 SNR'],
        ['5', 'Pfa / Pd', '★★★', '决定最小可检测 SNR 门限'],
        ['6', 'Fn / T0 → Pn_', '★★★', '噪声基底决定极限探测能力'],
        ['7', 'L (系统损耗)', '★★★', '分母项，每 3dB 损耗距离减少约 16%'],
    ]
)

heading(2, '4.2 侦察截获')
add_table(
    ['排序', '参数', '影响程度', '说明'],
    [
        ['1', 'mSensitivityDb (-70 dBm)', '★★★★★', '直接截获灵敏度门限'],
        ['2', 'mEsmFallbackBandwidthHz', '★★★★', '决定噪声功率基底'],
        ['3', 'mSnrScale (5000)', '★★★★', '控制测量误差幅度'],
        ['4', '天线增益 Gr', '★★★★', '公式中 Gt×Gr 直乘'],
        ['5', '入站雷达信号参数 Pt/PW/PRF', '★★★', '来自目标雷达，非本传感器控制'],
    ]
)

heading(2, '4.3 电子干扰')
add_table(
    ['排序', '参数', '影响程度', '说明'],
    [
        ['1', 'EqPower（干扰功率）', '★★★★★', '干扰方程分子，直接影响 J/S 比'],
        ['2', 'fj / Bj（频率与带宽）', '★★★★★', '决定频率重合度 BW_overlap/Bj'],
        ['3', 'antennaGain（干扰天线增益）', '★★★★', '干扰→雷达路径增益积'],
        ['4', 'KJC / KJI（雷达参数）', '★★★★', '干扰方程中的雷达脉压/积累增益'],
        ['5', 'EW 效果因子', '★★★★', '多层调制乘积效应可差数十分贝'],
    ]
)

# ============================================================
# 5. 文件索引
# ============================================================
heading(1, '5. 对应的源文件')
add_table(
    ['功能模块', '文件路径'],
    [
        ['雷达模型声明', 'externModel/sensor/source/RadarModel.hpp'],
        ['雷达模型实现', 'externModel/sensor/source/RadarModel.cpp'],
        ['雷达语法', 'externModel/sensor/grammar/sensor_cssc_01.ag'],
        ['雷达场景示例', 'externModel/sensor/data/sensor_scenario.txt'],
        ['侦察传感器声明', 'externModel/ReconSensor/source/WsfPassiveReconSensor.hpp'],
        ['侦察传感器实现', 'externModel/ReconSensor/source/WsfPassiveReconSensor.cpp'],
        ['侦察语法', 'externModel/ReconSensor/grammar/sensor_cssc_01.ag'],
        ['干扰武器声明', 'externModel/weapon/source/PhaserWeapon.hpp'],
        ['干扰武器实现', 'externModel/weapon/source/PhaserWeapon.cpp'],
        ['干扰语法', 'externModel/weapon/grammar/jammer_cssc_01.ag'],
        ['杀伤效果声明', 'externModel/weapon/source/PhaserLethality.hpp'],
        ['杀伤效果实现', 'externModel/weapon/source/PhaserLethality.cpp'],
        ['武器场景示例', 'externModel/weapon/data/weapon_scenario.txt'],
    ]
)

# ── 保存 ──
output_path = os.path.join(os.environ.get('USERPROFILE', 'C:\\'), 'Desktop', 'externModel_Parameter_Analysis.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
